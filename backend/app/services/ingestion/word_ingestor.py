from __future__ import annotations

import base64
import io

from app.models.document import DocumentInput
from app.services.ingestion.base import IngestorBase


class WordIngestor(IngestorBase):
    def ingest(self, doc: DocumentInput) -> str:
        from docx import Document

        # Content is base64-encoded docx bytes
        try:
            docx_bytes = base64.b64decode(doc.content)
            document = Document(io.BytesIO(docx_bytes))
        except Exception:
            # Might be a file path
            document = Document(doc.content)

        paragraphs = []
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        return "\n\n".join(paragraphs)
